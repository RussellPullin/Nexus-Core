#!/usr/bin/env python3
"""
Nexus Core — NDIS Document Template Preparation
================================================
Reads all .docx files from the Spring 2 Health policy library,
strips all PII (names, phone numbers, emails, ABN, org-specific branding),
replaces with Nexus token placeholders, writes clean .docx + PDF to
  data/forms/templates/library/<slug>/
and generates a manifest.json for each document.

Run from the nexus core project root:
  python3 scripts/prepare-ndis-templates.py
"""

import os
import re
import json
import shutil
import subprocess
import sys
from pathlib import Path
from copy import deepcopy

try:
    import docx
    from docx import Document
    from docx.oxml.ns import qn
except ImportError:
    print("ERROR: python-docx not installed. Run: pip3 install python-docx")
    sys.exit(1)

# ── Paths ──────────────────────────────────────────────────────────────────
SOURCE_DIR = Path(
    "/Users/pristinelifestylesolutions/Library/CloudStorage"
    "/OneDrive-PristineLifestyleSolutions/Spring 2 health"
    "/Updated policies S2H/Updated Word Docs S2H"
)
PROJECT_ROOT = Path(__file__).parent.parent
LIBRARY_OUT = PROJECT_ROOT / "data" / "forms" / "templates" / "library"
SOFFICE = shutil.which("soffice")

# ── PII replacement map ────────────────────────────────────────────────────
# Order matters — longer/more specific patterns first
PII_REPLACEMENTS = [
    # Specific people
    (re.compile(r"Kathleen\s+Hayes", re.I), "{org.primary_contact.name}"),
    (re.compile(r"Russell\s+Pullin", re.I), "{org.signatory.name}"),
    # Specific emails
    (re.compile(r"Russellpullin@gmail\.com", re.I), "{org.email}"),
    (re.compile(r"info@pristinelifestylesolutions\.com\.au", re.I), "{org.email}"),
    (re.compile(r"notify@comcare\.gov\.au", re.I), "notify@comcare.gov.au"),  # keep — regulator
    # Specific phone numbers
    (re.compile(r"\b0481\s*838\s*570\b"), "{org.phone}"),
    (re.compile(r"\b0468\s*404\s*845\b"), "{org.phone}"),
    # Generic phone patterns (AU) — only if not already tokenised
    (re.compile(r"\b(0[2-9]\d{8})\b"), "{org.phone}"),
    (re.compile(r"\+61\s?[2-9]\d{8}"), "{org.phone}"),
    # ABN — Spring 2 Health specific first, then generic 11-digit
    (re.compile(r"\b71\s*665\s*820\s*986\b"), "{org.abn}"),
    (re.compile(r"\bABN[:\s]+\d{2}\s*\d{3}\s*\d{3}\s*\d{3}\b", re.I), "ABN: {org.abn}"),
    # Org name variants — longest first
    (re.compile(r"Spring\s+2\s+Health\s+Pty\s+Ltd", re.I), "{org.legal_name}"),
    (re.compile(r"Spring\s+2\s+Health", re.I), "{org.name}"),
    (re.compile(r"Pristine\s+Lifestyle\s+Solutions", re.I), "{org.name}"),
    # Generic email pattern remaining after specific ones
    (re.compile(r"\b[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}\b"), "{org.email}"),
]

# Tokens Nexus already knows about (from templateTokens.js)
KNOWN_TOKENS = [
    "org.name", "org.legal_name", "org.abn", "org.email", "org.phone",
    "org.address", "org.website", "org.signatory.name", "org.signatory.role",
    "org.primary_contact.name", "org.primary_contact.role",
    "org.branding.logo_path", "today_long", "today",
    "participant.full_name", "participant.ndis_number", "participant.date_of_birth",
    "participant.email", "participant.phone", "participant.address",
    "participant.plan_start_date", "participant.plan_end_date",
    "staff.full_name", "staff.email", "staff.role", "staff.start_date",
]

# ── Document classification ────────────────────────────────────────────────
CATEGORY_RULES = [
    (re.compile(r"register$", re.I), "register"),
    (re.compile(r"(services?\s+agreement|service\s+agreement|services\s+agreement|support\s+coordination\s+services\s+agreement|sda.*sil.*agreement|sil.*agreement)", re.I), "contract"),
    (re.compile(r"(agreement|contract|engagement|declaration|consent)", re.I), "contract"),
    (re.compile(r"(policy$|policy\s+and|and\s+policy)", re.I), "policy"),
    (re.compile(r"(procedure|plan$|checklist|report|record|survey|booklet)", re.I), "procedure"),
    (re.compile(r"(form$|form\s+copy)", re.I), "form"),
    (re.compile(r"position\s+description|position description", re.I), "guide"),
]

PACK_RULES = [
    # participant onboarding
    (re.compile(r"(services?\s+agreement|client\s+intake|privacy\s+consent|client\s+induction|service\s+schedule|support\s+coordination\s+services)", re.I), "participant_onboarding"),
    # staff onboarding
    (re.compile(r"(staff\s+induction|worker\s+declaration|letter\s+of\s+engagement|contractor\s+agreement|position\s+description|reference\s+check|pre.?employment)", re.I), "staff_onboarding"),
    # compliance registers
    (re.compile(r"register$", re.I), "compliance_register"),
    # policies (staff read)
    (re.compile(r"policy$", re.I), "policy_library"),
]

NEEDS_SIGNATURE = {
    "services-agreement", "services-agreement-sil", "support-coordination-services-agreement",
    "privacy-consent-form", "sda-and-sil-collaboration-agreement",
    "worker-declarations", "conflict-of-interest-declaration",
    "independent-contractor-agreement", "letter-of-engagement-casual-employee",
    "client-intake-form", "advocacy-of-support-person-request-form",
    "change-of-supports", "exit-and-transition-form",
}

SIGNATURE_COUNTS = {
    "services-agreement": 2, "services-agreement-sil": 2,
    "support-coordination-services-agreement": 2, "sda-and-sil-collaboration-agreement": 2,
    "independent-contractor-agreement": 2, "letter-of-engagement-casual-employee": 2,
    "privacy-consent-form": 1, "worker-declarations": 1,
    "conflict-of-interest-declaration": 1, "client-intake-form": 1,
    "advocacy-of-support-person-request-form": 1, "change-of-supports": 1,
    "exit-and-transition-form": 1,
}

SKIP_DOCS = {
    # Internal business docs — not for other orgs
    "Business Plan.docx",
    "Core Module Answers.docx",
    "convert_word_to_pdf.py",
}

REGISTER_DOCS = {
    # These are auto-generated from DB — don't need docx templating
    "Complaints Register.docx",
    "Conflict of Interest Register.docx",
    "Continuous Improvement Register.docx",
    "Document Register.docx",
    "Emergency Test Register.docx",
    "Risk Assessed Role Register.docx",
    "Risk Register.docx",
    "Subject to a Significant Risk Factor Register.docx",
    "Training and Development Register.docx",
    "Waste Removal Records Register.docx",
    "Collection _ Storage of Medicine Register.docx",
}

# ── Helpers ────────────────────────────────────────────────────────────────

def slugify(name: str) -> str:
    name = re.sub(r"[_\s]+", "-", name.strip())
    name = re.sub(r"[^a-zA-Z0-9\-]", "", name)
    name = re.sub(r"-+", "-", name).strip("-").lower()
    return name

def classify(display_name: str) -> str:
    for pattern, cat in CATEGORY_RULES:
        if pattern.search(display_name):
            return cat
    return "policy"

def assign_pack(display_name: str) -> str:
    for pattern, pack in PACK_RULES:
        if pattern.search(display_name):
            return pack
    return "policy_library"

def clean_text(text: str) -> str:
    for pattern, replacement in PII_REPLACEMENTS:
        text = pattern.sub(replacement, text)
    return text

def collect_tokens(text: str) -> list:
    found = re.findall(r"\{([^}]+)\}", text)
    return sorted(set(t for t in found if t in KNOWN_TOKENS))

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"

def clean_xml_text_nodes(element):
    """Clean PII from all w:t (text) nodes within an XML element, including hyperlinks."""
    for t_node in element.iter(f"{{{W_NS}}}t"):
        if t_node.text:
            t_node.text = clean_text(t_node.text)

def clean_hyperlink_urls(element, doc):
    """Replace PII in hyperlink relationship targets (the actual URL)."""
    for hl in element.iter(f"{{{W_NS}}}hyperlink"):
        r_id = hl.get(f"{{{R_NS}}}id")
        if r_id:
            try:
                rel = doc.part.rels.get(r_id)
                if rel and rel.target_ref:
                    cleaned = clean_text(rel.target_ref)
                    if cleaned != rel.target_ref:
                        rel._target = cleaned
            except Exception:
                pass

def process_paragraph(para):
    # Clean all text nodes including hyperlinks at XML level
    clean_xml_text_nodes(para._element)
    clean_hyperlink_urls(para._element, para._p.getroottree().getroot() if hasattr(para._p, 'getroottree') else None)

def process_paragraph_with_doc(para, doc):
    clean_xml_text_nodes(para._element)
    clean_hyperlink_urls(para._element, doc)

def process_table(table, doc):
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                process_paragraph_with_doc(para, doc)
            for tbl in cell.tables:
                process_table(tbl, doc)

def process_document(doc_path: Path) -> tuple[Document, list]:
    doc = Document(str(doc_path))

    for para in doc.paragraphs:
        process_paragraph_with_doc(para, doc)

    for table in doc.tables:
        process_table(table, doc)

    # Process headers and footers
    for section in doc.sections:
        for hf in [section.header, section.footer]:
            if hf:
                for para in hf.paragraphs:
                    process_paragraph_with_doc(para, doc)
                for table in hf.tables:
                    process_table(table, doc)

    # Collect all tokens now present
    all_text = " ".join(
        "".join(r.text for r in p.runs) for p in doc.paragraphs
    )
    for section in doc.sections:
        for hf in [section.header, section.footer]:
            if hf:
                all_text += " " + " ".join("".join(r.text for r in p.runs) for p in hf.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_text += " " + cell.text

    tokens = collect_tokens(all_text)
    return doc, tokens

def convert_to_pdf(docx_path: Path, out_dir: Path) -> Path:
    if not SOFFICE:
        print("  WARNING: soffice not found — skipping PDF conversion")
        return None
    result = subprocess.run(
        [SOFFICE, "--headless", "--convert-to", "pdf", "--outdir", str(out_dir), str(docx_path)],
        capture_output=True, text=True, timeout=60
    )
    pdf_name = docx_path.stem + ".pdf"
    pdf_path = out_dir / pdf_name
    if pdf_path.exists():
        return pdf_path
    print(f"  WARNING: PDF conversion failed for {docx_path.name}: {result.stderr[:200]}")
    return None

def make_manifest(slug, display_name, category, pack, tokens, has_sig, sig_count, has_pdf):
    template_file = "template.pdf" if has_pdf else "template.docx"
    engine = "pdf-acroform" if has_pdf else "docxtemplater"
    return {
        "slug": slug,
        "display_name": display_name,
        "category": category,
        "form_type": category,
        "engine": engine,
        "version": "1.0.0",
        "template_file": template_file,
        "pack": pack,
        "placeholders": tokens,
        "required_signer_role": "participant" if has_sig else None,
        "signature_count": sig_count,
        "renewal_days": 365 if category == "contract" else None,
        "is_active": True,
    }

# ── Main ───────────────────────────────────────────────────────────────────

def main():
    if not SOURCE_DIR.exists():
        print(f"ERROR: Source directory not found:\n  {SOURCE_DIR}")
        sys.exit(1)

    LIBRARY_OUT.mkdir(parents=True, exist_ok=True)

    docs = sorted(SOURCE_DIR.glob("*.docx"))
    print(f"\nNexus Core — NDIS Template Preparation")
    print(f"Source: {SOURCE_DIR}")
    print(f"Output: {LIBRARY_OUT}")
    print(f"Found:  {len(docs)} .docx files\n")

    results = {"processed": 0, "skipped": 0, "errors": 0, "by_category": {}, "by_pack": {}}
    manifest_list = []

    for doc_path in docs:
        if doc_path.name in SKIP_DOCS:
            print(f"  SKIP  {doc_path.name} (internal doc)")
            results["skipped"] += 1
            continue

        is_register = doc_path.name in REGISTER_DOCS
        display_name = doc_path.stem.replace(" copy", "").strip()
        slug = slugify(display_name)
        category = "register" if is_register else classify(display_name)
        pack = assign_pack(display_name)
        has_sig = slug in NEEDS_SIGNATURE
        sig_count = SIGNATURE_COUNTS.get(slug, 1 if has_sig else 0)

        out_dir = LIBRARY_OUT / slug
        out_dir.mkdir(exist_ok=True)

        print(f"  PROC  {doc_path.name}")

        try:
            if is_register:
                # Registers: copy as-is for reference, mark as db-generated
                shutil.copy2(doc_path, out_dir / "template.docx")
                tokens = []
                pdf_path = convert_to_pdf(out_dir / "template.docx", out_dir)
            else:
                doc, tokens = process_document(doc_path)
                clean_docx = out_dir / "template.docx"
                doc.save(str(clean_docx))
                pdf_path = convert_to_pdf(clean_docx, out_dir)

            has_pdf = pdf_path is not None and pdf_path.exists()

            manifest = make_manifest(slug, display_name, category, pack, tokens, has_sig, sig_count, has_pdf)
            manifest_json = out_dir / "manifest.json"
            manifest_json.write_text(json.dumps(manifest, indent=2, default=str))

            manifest_list.append(manifest)
            results["processed"] += 1
            results["by_category"][category] = results["by_category"].get(category, 0) + 1
            results["by_pack"][pack] = results["by_pack"].get(pack, 0) + 1

            pdf_note = "✓ PDF" if has_pdf else "✗ no PDF"
            sig_note = f"sig×{sig_count}" if has_sig else ""
            token_note = f"tokens: {','.join(tokens)}" if tokens else "no tokens"
            print(f"         → [{category}] {pdf_note} {sig_note}  {token_note}")

        except Exception as e:
            print(f"  ERROR {doc_path.name}: {e}")
            results["errors"] += 1

    # Write full catalogue
    catalogue_path = LIBRARY_OUT / "_catalogue.json"
    catalogue_path.write_text(json.dumps(manifest_list, indent=2, default=str))

    print(f"\n{'='*60}")
    print(f"Done: {results['processed']} processed, {results['skipped']} skipped, {results['errors']} errors")
    print(f"\nBy category:")
    for cat, n in sorted(results["by_category"].items()):
        print(f"  {cat:20} {n}")
    print(f"\nBy pack:")
    for pack, n in sorted(results["by_pack"].items()):
        print(f"  {pack:30} {n}")
    print(f"\nLibrary: {LIBRARY_OUT}")
    print(f"Catalogue: {catalogue_path}")

if __name__ == "__main__":
    main()
